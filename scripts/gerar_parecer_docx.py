#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Gera versão Word (.docx) bem layoutizada do PARECER_IRRF_VIGILANCIA_LIMPEZA_2026-04-17.md
e salva em C:\\Users\\Avell\\Downloads\\.
"""
import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

SRC = r"C:\Users\Avell\OneDrive\Área de Trabalho\Montana_Seg_Conciliacao\app_unificado\PARECER_IRRF_VIGILANCIA_LIMPEZA_2026-04-17.md"
DST = r"C:\Users\Avell\Downloads\PARECER_IRRF_VIGILANCIA_LIMPEZA_2026-04-17.docx"

# ── cores institucionais ──────────────────────────────────────────
C_PRIMARY = RGBColor(0x0B, 0x3D, 0x91)   # azul escuro
C_SECONDARY = RGBColor(0x22, 0x55, 0xA4)
C_ACCENT  = RGBColor(0xC1, 0x27, 0x27)   # vermelho alerta
C_TEXT    = RGBColor(0x1F, 0x1F, 0x1F)
C_MUTED   = RGBColor(0x55, 0x55, 0x55)
C_TABLE_HDR_BG = "0B3D91"
C_TABLE_ALT_BG = "F1F5FB"

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def set_cell_borders(cell, color="BFBFBF", sz="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:color'), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)

def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:color'), '0B3D91')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_page_number_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Grupo Montana — Parecer Técnico-Jurídico IRRF — Página ")
    run.font.size = Pt(8)
    run.font.color.rgb = C_MUTED
    # page field
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = "1"
    r.append(t)
    fld.append(r)
    p._p.append(fld)

def render_inline(paragraph, text, base_color=None, base_size=None):
    """Parses **bold**, *italic*, `code`, ~~strike~~ and emits formatted runs."""
    # split preserving tokens
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|~~[^~]+~~)'
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        if part.startswith('**') and part.endswith('**'):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run.text = part[1:-1]
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run.text = part[1:-1]
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        elif part.startswith('~~') and part.endswith('~~'):
            run.text = part[2:-2]
            run.font.strike = True
        else:
            # remove markdown escape artifacts
            run.text = part
        if base_color:
            run.font.color.rgb = base_color
        if base_size:
            run.font.size = base_size

def add_heading(doc, text, level):
    # level 1 = ##, level 2 = ###, level 3 = ####
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12 if level == 2 else 8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16); run.font.color.rgb = C_PRIMARY
    elif level == 2:
        run.font.size = Pt(13); run.font.color.rgb = C_SECONDARY
    else:
        run.font.size = Pt(11); run.font.color.rgb = C_TEXT
    return p

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = C_PRIMARY

def add_cover_metadata(doc, meta_lines):
    """meta_lines: list of (label, value) tuples"""
    tbl = doc.add_table(rows=len(meta_lines), cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Cm(4.5)
    tbl.columns[1].width = Cm(11.5)
    for i, (label, value) in enumerate(meta_lines):
        c1 = tbl.cell(i, 0)
        c1.width = Cm(4.5)
        c1.text = ''
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(label)
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = C_PRIMARY
        c2 = tbl.cell(i, 1)
        c2.width = Cm(11.5)
        c2.text = ''
        p2 = c2.paragraphs[0]
        render_inline(p2, value, base_color=C_TEXT, base_size=Pt(10))
        set_cell_borders(c1, color="FFFFFF")
        set_cell_borders(c2, color="FFFFFF")
    doc.add_paragraph()

def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # left border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), '0B3D91')
    pBdr.append(left)
    pPr.append(pBdr)
    # shading
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F1F5FB')
    pPr.append(shd)
    render_inline(p, text, base_color=C_TEXT, base_size=Pt(10))
    for run in p.runs:
        run.italic = True

def add_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    render_inline(p, text, base_color=C_TEXT, base_size=Pt(11))

def add_list_item(doc, text, level=0, numbered=False, idx=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7 + level * 0.6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    bullet = f"{idx}. " if numbered and idx else "• "
    r = p.add_run(bullet)
    r.bold = True
    r.font.color.rgb = C_PRIMARY
    r.font.size = Pt(11)
    render_inline(p, text, base_color=C_TEXT, base_size=Pt(11))

def strip_md_cell(cell):
    return cell.strip().strip('|').strip()

def add_table_from_md(doc, lines):
    """lines: list of pipe-delimited rows (incl header + separator + body)."""
    rows = []
    alignments = []
    for line in lines:
        if re.match(r'^\s*\|?\s*:?\-+:?\s*(\|\s*:?\-+:?\s*)+\|?\s*$', line):
            # alignment row
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            for c in cells:
                if c.startswith(':') and c.endswith(':'):
                    alignments.append(WD_ALIGN_PARAGRAPH.CENTER)
                elif c.endswith(':'):
                    alignments.append(WD_ALIGN_PARAGRAPH.RIGHT)
                else:
                    alignments.append(WD_ALIGN_PARAGRAPH.LEFT)
            continue
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    while len(alignments) < ncols:
        alignments.append(WD_ALIGN_PARAGRAPH.LEFT)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = tbl.cell(i, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = alignments[j]
            txt = row[j] if j < len(row) else ''
            if i == 0:
                render_inline(p, txt, base_color=RGBColor(0xFF,0xFF,0xFF), base_size=Pt(10))
                for r in p.runs:
                    r.bold = True
                set_cell_bg(cell, C_TABLE_HDR_BG)
            else:
                render_inline(p, txt, base_color=C_TEXT, base_size=Pt(9.5))
                if i % 2 == 0:
                    set_cell_bg(cell, C_TABLE_ALT_BG)
            set_cell_borders(cell, color="BFBFBF")
    # spacing after
    doc.add_paragraph()

# ── PARSE MARKDOWN ─────────────────────────────────────────────────
def main():
    with open(SRC, 'r', encoding='utf-8') as f:
        md = f.read()

    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Fonte padrão
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = C_TEXT

    add_page_number_footer(doc)

    # Header on first page
    # Cover block
    cover_p = doc.add_paragraph()
    cover_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_p.paragraph_format.space_after = Pt(4)
    r = cover_p.add_run("GRUPO MONTANA")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = C_MUTED

    add_title(doc, "PARECER TÉCNICO-JURÍDICO")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    rr = subtitle.add_run("IRRF 1,20% — Vigilância, Segurança, Limpeza, Conservação e Locação de Mão de Obra")
    rr.italic = True
    rr.font.size = Pt(12)
    rr.font.color.rgb = C_SECONDARY

    add_hr(doc)

    # Parse the rest of the markdown
    lines = md.splitlines()

    # Skip the title line "# PARECER..."
    i = 0
    while i < len(lines) and not lines[i].startswith('**Assunto:'):
        i += 1

    # Collect cover metadata lines
    cover_data = []
    while i < len(lines) and (lines[i].startswith('**') or lines[i].strip() == ''):
        line = lines[i].strip()
        if line.startswith('**') and ':**' in line:
            m = re.match(r'\*\*([^*]+):\*\*\s*(.+)', line)
            if m:
                cover_data.append((m.group(1).strip() + ':', m.group(2).strip()))
        i += 1
    if cover_data:
        add_cover_metadata(doc, cover_data)

    add_hr(doc)

    # Process body
    buf_table = []
    in_table = False

    def flush_table():
        nonlocal buf_table, in_table
        if buf_table:
            add_table_from_md(doc, buf_table)
            buf_table = []
        in_table = False

    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.rstrip()

        # Table detection
        if stripped.startswith('|') and '|' in stripped[1:]:
            in_table = True
            buf_table.append(stripped)
            i += 1
            continue
        elif in_table:
            flush_table()

        # Skip blank
        if not stripped.strip():
            i += 1
            continue

        # Horizontal rule
        if stripped.strip() == '---':
            add_hr(doc)
            i += 1
            continue

        # Headings
        if stripped.startswith('#### '):
            add_heading(doc, stripped[5:].strip(), 3)
            i += 1
            continue
        if stripped.startswith('### '):
            add_heading(doc, stripped[4:].strip(), 2)
            i += 1
            continue
        if stripped.startswith('## '):
            add_heading(doc, stripped[3:].strip(), 1)
            i += 1
            continue
        if stripped.startswith('# '):
            # skip (already rendered as title)
            i += 1
            continue

        # Blockquote (possibly multi-line)
        if stripped.startswith('> '):
            block = []
            while i < n and lines[i].startswith('>'):
                ln = lines[i][1:].lstrip()
                if ln == '':
                    block.append('')
                else:
                    block.append(ln)
                i += 1
            # Join preserving paragraph breaks
            text = '\n'.join(block).strip()
            # split by blank lines into sub-paragraphs
            subs = re.split(r'\n\s*\n', text)
            for s in subs:
                add_quote(doc, s.replace('\n', ' '))
            continue

        # Ordered list
        m_ol = re.match(r'^(\s*)(\d+)\.\s+(.+)', line)
        if m_ol:
            indent_spaces = len(m_ol.group(1))
            level = indent_spaces // 2
            add_list_item(doc, m_ol.group(3), level=level, numbered=True, idx=m_ol.group(2))
            i += 1
            continue

        # Unordered list
        m_ul = re.match(r'^(\s*)[-*]\s+(.+)', line)
        if m_ul:
            indent_spaces = len(m_ul.group(1))
            level = indent_spaces // 2
            add_list_item(doc, m_ul.group(2), level=level, numbered=False)
            i += 1
            continue

        # Default paragraph (may span continuation lines handled by each line)
        # Italic closing paragraph
        if stripped.startswith('*') and stripped.endswith('*') and stripped.count('*') >= 2 and not stripped.startswith('**'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            r = p.add_run(stripped.strip('*'))
            r.italic = True
            r.font.size = Pt(9.5)
            r.font.color.rgb = C_MUTED
            i += 1
            continue

        add_paragraph(doc, stripped)
        i += 1

    if in_table:
        flush_table()

    doc.save(DST)
    print(f"[OK] Arquivo gerado: {DST}")
    print(f"     Tamanho: {os.path.getsize(DST):,} bytes")

if __name__ == '__main__':
    main()
